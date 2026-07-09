# -*- coding: utf-8 -*-
"""
ingest.py — قراءة الملفات من documents/ وتقطيعها تقطيعًا ذكيًا (فقرات + صفحات)
ثم تحويلها لمتجهات عبر bge-m3 وتخزينها في ChromaDB.
الاستخدام: python3 ingest.py
"""
import os
import re
import sys
from pathlib import Path

import chromadb
import ollama

BASE_DIR = Path(__file__).parent
DOCS_DIR = BASE_DIR / "documents"
DB_DIR = BASE_DIR / "chroma_db"

EMBED_MODEL = os.environ.get("RAG_EMBED_MODEL", "bge-m3")
COLLECTION = "documents"
CHUNK_WORDS = int(os.environ.get("RAG_CHUNK_WORDS", "300"))     # الحجم المستهدف (كلمات)
OVERLAP_WORDS = int(os.environ.get("RAG_OVERLAP_WORDS", "75"))  # التداخل (كلمات)
MIN_CHUNK_WORDS = 25   # أصغر مقطع مقبول
BATCH_SIZE = 16

# ---------- OCR للمستندات الممسوحة ضوئيًا/الصور (عبر نموذج رؤية محلي في Ollama) ----------
OCR_ENABLED = os.environ.get("RAG_OCR", "1").lower() not in ("0", "false", "no")
OCR_MODEL = os.environ.get("RAG_OCR_MODEL", "qwen2.5vl:3b")  # خفيف لـ MacBook Air 16GB؛ للجهاز القوي qwen2.5vl:32b
OCR_MIN_CHARS = int(os.environ.get("RAG_OCR_MIN_CHARS", "20"))  # صفحة نصها أقل من ذلك = ممسوحة
OCR_DPI = int(os.environ.get("RAG_OCR_DPI", "200"))
OCR_PROMPT = ("استخرج كل النص الظاهر في هذه الصورة حرفيًا وبدقة تامة باللغة العربية "
              "(وأي نص إنجليزي إن وُجد)، مع الحفاظ على ترتيب الأسطر والفقرات والأرقام كما هي. "
              "أخرج النص المستخرج فقط دون أي شرح أو تعليق أو مقدمات.")


def ocr_image(png_bytes):
    """OCR لصورة (bytes بصيغة PNG) عبر نموذج الرؤية المحلي. يُرجع النص أو '' عند الفشل."""
    try:
        resp = ollama.chat(
            model=OCR_MODEL,
            messages=[{"role": "user", "content": OCR_PROMPT, "images": [png_bytes]}],
            options={"temperature": 0},
        )
        msg = resp["message"] if isinstance(resp, dict) else resp.message
        return (msg["content"] if isinstance(msg, dict) else msg.content).strip()
    except Exception as e:
        print(f"⚠️ تعذّر تشغيل OCR (النموذج {OCR_MODEL}): {e}")
        return ""

# فواصل الجمل العربية واللاتينية
_SENT_SPLIT = re.compile(r"(?<=[.!؟?؛…])\s+")
# عنصر قائمة مرقّمة/منقّطة (لحماية القوائم من الانقسام بين مقطعين)
_LIST_ITEM = re.compile(r"^\s*(?:[\d٠-٩]+[.)\-–]|[•▪‣●○*])\s+")


# ---------- قراءة الملفات: تُرجع قائمة (نص_فقرة، رقم_صفحة أو None) ----------

HEAD_MARK = "\x00HEADING\x00"  # علامة داخلية: عنوان قسم رئيسي

# كلمات عربية قصيرة تقف وحدها — لا تُلصق بجاراتها
_STANDALONE = {"في", "من", "عن", "ما", "لا", "أو", "أن", "إن", "لم", "لن", "له",
               "هو", "هي", "ثم", "قد", "كل", "أي", "يا", "إذ", "بل", "لو", "ال",
               "بن", "أب", "ذا", "ذي", "مع", "إذا"}
_AR_WORD = re.compile(r"[ء-يٰ-ۓ]+")


def _normalize_arabic(text: str) -> str:
    """إصلاح شامل لنص PDF العربي:
    1) NFKC: تحويل أشكال العرض (ﺗ، ﻻ...) إلى حروف قياسية
    2) قلب اللام-ألف المعكوسة (األ → الأ، اال → الا) — خلل معروف في بعض المستخرجات
    3) «ال» المستقلة → «لا» (نفي مقلوب)
    4) إزالة الكشيدة ثم لصق شظايا الكلمات"""
    import unicodedata
    text = unicodedata.normalize("NFKC", text)
    for wrong, right in (("األ", "الأ"), ("اال", "الا"), ("اإل", "الإ"), ("اآل", "الآ")):
        text = text.replace(wrong, right)
    text = re.sub(r"(?<![ء-ي])ال(?=\s)", "لا", text)
    return _fix_fragments(text.replace("ـ", ""))


def _fix_fragments(text: str) -> str:
    """لصق شظايا الكلمات العربية الناتجة عن بعض مستخرجات PDF: «تنطب ق» → «تنطبق»."""
    def frag(w):
        return (len(w) <= 2 and w not in _STANDALONE
                and bool(_AR_WORD.fullmatch(w)))

    out_lines = []
    for line in text.split("\n"):
        out = []
        for w in line.split(" "):
            if (out and w and (frag(w) or frag(out[-1]))
                    and _AR_WORD.fullmatch(w) and _AR_WORD.fullmatch(out[-1])):
                out[-1] += w
            else:
                out.append(w)
        out_lines.append(" ".join(out))
    return "\n".join(out_lines)


def _is_scanned(text):
    """صفحة بلا طبقة نص فعلية (ممسوحة ضوئيًا) = عدد الأحرف غير الفراغية أقل من الحد."""
    return len(re.sub(r"\s", "", text or "")) < OCR_MIN_CHARS


def read_pdf(path: Path):
    # PyMuPDF أدق بكثير مع العربية والأرقام (اتجاه RTL)؛ pypdf احتياط
    try:
        import fitz
        with fitz.open(str(path)) as doc:
            pages = []
            for i, page in enumerate(doc):
                text = page.get_text()
                # صفحة ممسوحة ضوئيًا (بلا نص) → OCR عبر نموذج الرؤية المحلي
                if OCR_ENABLED and _is_scanned(text):
                    try:
                        pix = page.get_pixmap(dpi=OCR_DPI)
                        ocr = ocr_image(pix.tobytes("png"))
                        if ocr:
                            text = ocr
                    except Exception as e:
                        print(f"⚠️ تعذّر OCR للصفحة {i + 1}: {e}")
                pages.append((i + 1, text))
    except ImportError:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        pages = [(i + 1, pg.extract_text() or "") for i, pg in enumerate(reader.pages)]

    parts = []
    for page_no, text in pages:
        text = _normalize_arabic(text)
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        content = [l for l in lines if not re.fullmatch(r"[\d٠-٩]+", l)]
        # صفحة غلاف قسم: سطر أو سطران قصيران فقط (عدا رقم الصفحة) → عنوان رئيسي
        if 0 < len(content) <= 2 and sum(len(l.split()) for l in content) <= 8:
            title = re.sub(r"\s+", " ", " ".join(content).replace("\t", " ")).strip()
            parts.append((HEAD_MARK + title, page_no))
            continue
        for para in re.split(r"\n\s*\n", "\n".join(lines)):
            para = para.strip()
            if para:
                parts.append((para, page_no))
    return parts


def read_docx(path: Path):
    from docx import Document
    doc = Document(str(path))
    parts = [(p.text.strip(), None) for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            line = " | ".join(cell.text.strip() for cell in row.cells)
            if line.strip(" |"):
                parts.append((line, None))
    return parts


def read_text(path: Path):
    raw = None
    for enc in ("utf-8", "utf-8-sig", "cp1256"):
        try:
            raw = path.read_text(encoding=enc)
            break
        except UnicodeDecodeError:
            continue
    if raw is None:
        raw = path.read_text(encoding="utf-8", errors="replace")
    return [(p.strip(), None) for p in re.split(r"\n\s*\n", raw) if p.strip()]


def read_image(path: Path):
    """OCR لملف صورة (ممسوح ضوئيًا) عبر نموذج الرؤية المحلي."""
    if not OCR_ENABLED:
        return []
    text = _normalize_arabic(ocr_image(path.read_bytes()))
    return [(p.strip(), None) for p in re.split(r"\n\s*\n", text) if p.strip()]


READERS = {".pdf": read_pdf, ".docx": read_docx, ".txt": read_text, ".md": read_text,
           ".png": read_image, ".jpg": read_image, ".jpeg": read_image}


# ---------- التقطيع الذكي (بحدود الجمل + اكتشاف العناوين) ----------

_MAJOR_HEAD = re.compile(r"^\s*[\d٠-٩]+(?:[.][\d٠-٩]+)?\s*[.)\-–]?\s+\S")


def _is_heading(text: str) -> bool:
    """عنوان قسم رئيسي احتياطي (لملفات بلا صفحات غلاف): سطر قصير يبدأ بترقيم
    من مستوى أو مستويين (مثل: 3. الأهداف أو 4.3 سياسة كذا) دون رقم في نهايته.
    العناوين الفرعية القصيرة (كأسماء المبادئ) تبقى داخل النص ولا تُنشئ قسمًا."""
    words = text.split()
    if not (2 <= len(words) <= 8) or len(text) > 80:
        return False
    if re.search(r"[.؟!؛,،:]\s*$", text) or re.search(r"[\d٠-٩]+\s*$", text):
        return False
    return bool(_MAJOR_HEAD.match(text))


def _sentences(text: str):
    """تقسيم نص إلى جمل (يدعم علامات الترقيم العربية)."""
    return [s.strip() for s in _SENT_SPLIT.split(text) if s.strip()]


def chunk_paragraphs(paragraphs):
    """
    تجميع الجمل في مقاطع ~CHUNK_WORDS كلمة بتداخل جُملي ~OVERLAP_WORDS،
    مع إلحاق عنوان القسم الحالي بكل مقطع وتتبّع نطاق الصفحات.
    يُرجع قائمة dict: {text, pages, heading}
    """
    # وحدات = جمل مع صفحتها وعنوان قسمها وهل هي ضمن قائمة
    units = []  # (words, page, heading, islist)
    heading = ""
    for text, page in paragraphs:
        if text.startswith(HEAD_MARK):  # عنوان رئيسي من صفحة غلاف
            heading = text[len(HEAD_MARK):]
            continue
        if _is_heading(text):
            heading = re.sub(r"\s+", " ", text).strip()
            continue
        islist = bool(_LIST_ITEM.match(text)) or bool(
            re.search(r"\n\s*(?:[\d٠-٩]+[.)\-–]|[•▪‣●○*])\s+", text))
        for sent in _sentences(text):
            w = sent.split()
            while len(w) > CHUNK_WORDS:  # جملة أطول من الحد → تُقسم قسرًا
                units.append((w[:CHUNK_WORDS], page, heading, islist))
                w = w[CHUNK_WORDS:]
            if w:
                units.append((w, page, heading, islist))

    def make(buf):
        head = next((h for _, _, h, _ in buf if h), "")
        text = " ".join(" ".join(w) for w, _, _, _ in buf)
        if head:
            text = f"({head}) {text}"
        pages = sorted({p for _, p, _, _ in buf if p is not None})
        return {"text": text, "heading": head,
                "pages": f"{pages[0]}-{pages[-1]}" if pages else ""}

    chunks, buf, count, fresh = [], [], 0, 0
    cur_head = None

    def close_section():
        """إغلاق المقطع عند نهاية قسم — لا مقطع ولا تداخل يعبر حدود قسمين."""
        nonlocal buf, count, fresh
        if buf and fresh > 0:
            if fresh >= 10 or not chunks:
                chunks.append(make(buf))
            elif chunks[-1]["heading"] == (buf[0][2] or ""):
                chunks[-1]["text"] += " " + " ".join(" ".join(w) for w, _, _, _ in buf)
        buf, count, fresh = [], 0, 0

    for idx, u in enumerate(units):
        if buf and u[2] != cur_head:
            close_section()  # بداية قسم جديد
        cur_head = u[2]
        buf.append(u)
        count += len(u[0])
        fresh += len(u[0])
        # لا نقطع في منتصف قائمة: نتجاوز الحد حتى +200 كلمة إذا كانت القائمة مستمرة
        nxt = units[idx + 1] if idx + 1 < len(units) else None
        mid_list = u[3] and nxt is not None and nxt[3] and nxt[2] == u[2]
        if count >= CHUNK_WORDS + 200 or (count >= CHUNK_WORDS and not mid_list):
            chunks.append(make(buf))
            # التداخل داخل القسم نفسه فقط: آخر جمل بمجموع ~OVERLAP_WORDS
            keep, kc = [], 0
            for unit in reversed(buf):
                keep.insert(0, unit)
                kc += len(unit[0])
                if kc >= OVERLAP_WORDS:
                    break
            buf, count, fresh = keep, kc, 0
    close_section()
    return chunks


# ---------- التضمين والتخزين ----------

def embed(texts):
    return ollama.embed(model=EMBED_MODEL, input=texts)["embeddings"]


def main():
    if not DOCS_DIR.exists():
        sys.exit(f"المجلد غير موجود: {DOCS_DIR}")

    files = sorted(p for p in DOCS_DIR.rglob("*") if p.suffix.lower() in READERS)
    if not files:
        sys.exit("لا توجد ملفات مدعومة (pdf, docx, txt, md) داخل documents/")

    client = chromadb.PersistentClient(path=str(DB_DIR))
    col = client.get_or_create_collection(COLLECTION, metadata={"hnsw:space": "cosine"})

    total = 0
    for path in files:
        name = path.name
        print(f"⏳ معالجة: {name}")
        try:
            paragraphs = READERS[path.suffix.lower()](path)
        except Exception as e:
            print(f"  ⚠️ تعذّرت قراءة {name}: {e}")
            continue
        chunks = chunk_paragraphs(paragraphs)
        if not chunks:
            print(f"  ⚠️ {name} فارغ — تم تخطيه")
            continue

        # حذف الإدخالات القديمة لهذا الملف لتفادي التكرار عند إعادة الفهرسة
        col.delete(where={"source": name})

        for i in range(0, len(chunks), BATCH_SIZE):
            batch = chunks[i:i + BATCH_SIZE]
            vectors = embed([c["text"] for c in batch])
            col.add(
                ids=[f"{name}::{i + j}" for j in range(len(batch))],
                embeddings=vectors,
                documents=[c["text"] for c in batch],
                metadatas=[{"source": name, "pages": c["pages"], "chunk": i + j,
                            "heading": c.get("heading", "")}
                           for j, c in enumerate(batch)],
            )
            print(f"  … {min(i + BATCH_SIZE, len(chunks))}/{len(chunks)}", end="\r")
        total += len(chunks)
        print(f"  ✅ {len(chunks)} مقطع")

    # حذف إدخالات الملفات التي أُزيلت من documents/
    current = {p.name for p in files}
    existing = col.get(include=["metadatas"])
    orphans = {m["source"] for m in existing["metadatas"]} - current
    for src in orphans:
        col.delete(where={"source": src})
        print(f"🗑️ حُذف من القاعدة (الملف لم يعد موجودًا): {src}")

    print(f"\nتمت الفهرسة: {len(files)} ملف، {total} مقطع. القاعدة: {DB_DIR}")


if __name__ == "__main__":
    main()
